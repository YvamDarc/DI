import os
import json
import time
from typing import List, Dict, Any
from pathlib import Path

import streamlit as st
import pandas as pd
from docx import Document
import re

# -----------------------
# Helpers & constants
# -----------------------

BASE_DIR = Path(__file__).parent
FORMS_DIR = BASE_DIR / "forms"
RESP_DIR = BASE_DIR / "responses"
UPLOADS_DIR = BASE_DIR / "uploads"
WORD_DIR = BASE_DIR / "word_exports"
DRAFTS_DIR = BASE_DIR / "drafts"

FORMS_DIR.mkdir(exist_ok=True)
RESP_DIR.mkdir(exist_ok=True)
UPLOADS_DIR.mkdir(exist_ok=True)
WORD_DIR.mkdir(exist_ok=True)
DRAFTS_DIR.mkdir(exist_ok=True)

ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "admin")

SUPPORTED_TYPES = ["oui_non", "texte", "checkbox_multi", "fichier"]

def load_form(client_id: str) -> List[Dict[str, Any]]:
    path = FORMS_DIR / f"{client_id}.json"
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8"))

def save_form(client_id: str, questions: List[Dict[str, Any]]):
    path = FORMS_DIR / f"{client_id}.json"
    path.write_text(json.dumps(questions, ensure_ascii=False, indent=2), encoding="utf-8")

def load_draft(client_id: str) -> Dict[str, Any]:
    path = DRAFTS_DIR / f"{client_id}.json"
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}

def save_draft(client_id: str, draft: Dict[str, Any]):
    path = DRAFTS_DIR / f"{client_id}.json"
    path.write_text(json.dumps(draft, ensure_ascii=False, indent=2), encoding="utf-8")

def export_word(questions: List[Dict[str, Any]], client_id: str) -> Path:
    doc = Document()
    doc.add_heading(f"Formulaire – {client_id}", 0)
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"Q{i}. {q['question']}")
    out = WORD_DIR / f"{client_id}.docx"
    doc.save(out)
    return out

def append_responses_csv(client_id: str, data_rows: List[Dict[str, Any]]):
    out = RESP_DIR / f"{client_id}.csv"
    df_new = pd.DataFrame(data_rows)
    if out.exists():
        df_old = pd.read_csv(out)
        df = pd.concat([df_old, df_new], ignore_index=True)
    else:
        df = df_new
    df.to_csv(out, index=False)

def normalize_options(opt_str: str) -> List[str]:
    if not isinstance(opt_str, str):
        return []
    parts = [o.strip() for o in str(opt_str).split(";")]
    return [p for p in parts if p]

# -----------------------
# FEC analysis (401/411/471) + seuil de matérialité
# -----------------------

def _colmap(df: pd.DataFrame) -> dict:
    return {c.lower(): c for c in df.columns}

def _getcol(df: pd.DataFrame, *candidates: str):
    cmap = _colmap(df)
    for cand in candidates:
        if cand.lower() in cmap:
            return cmap[cand.lower()]
    return None

def _starts_with_series(s: pd.Series, prefixes):
    if isinstance(prefixes, str):
        prefixes = (prefixes,)
    return s.astype(str).str.startswith(prefixes, na=False)

def _empty_series(s: pd.Series):
    return s.isna() | (s.astype(str).str.strip() == "")

def _make_amount_series(df: pd.DataFrame, cDebit: str, cCredit: str):
    if cDebit and (cDebit in df.columns) and cCredit and (cCredit in df.columns):
        d = pd.to_numeric(df[cDebit].str.replace(",", ".", regex=False), errors="coerce").fillna(0.0)
        c = pd.to_numeric(df[cCredit].str.replace(",", ".", regex=False), errors="coerce").fillna(0.0)
        return (d - c).abs()
    elif cDebit and (cDebit in df.columns):
        return pd.to_numeric(df[cDebit].str.replace(",", ".", regex=False), errors="coerce").abs().fillna(0.0)
    elif cCredit and (cCredit in df.columns):
        return pd.to_numeric(df[cCredit].str.replace(",", ".", regex=False), errors="coerce").abs().fillna(0.0)
    else:
        return pd.Series(0.0, index=df.index)

def read_fec_autodetect(uploaded_file):
    # 1) CSV auto-sep
    try:
        df = pd.read_csv(uploaded_file, sep=None, engine="python", dtype=str)
        return df
    except Exception:
        uploaded_file.seek(0)
    # 2) CSV ; séparateur
    try:
        df = pd.read_csv(uploaded_file, sep=";", dtype=str)
        return df
    except Exception:
        uploaded_file.seek(0)
    # 3) Excel
    return pd.read_excel(uploaded_file)

def generate_questions_401_411_471(
    df: pd.DataFrame,
    aging_days: int = 90,
    top_n_per_bucket: int = 20,
    amount_threshold: float = 0.0
) -> List[Dict[str, Any]]:
    q: List[Dict[str, Any]] = []

    # Colonnes
    cCompteNum = _getcol(df, "CompteNum")
    cCompteLib = _getcol(df, "CompteLib")
    cCompAuxNum = _getcol(df, "CompAuxNum")
    cCompAuxLib = _getcol(df, "CompAuxLib")
    cEcrDate   = _getcol(df, "EcritureDate", "DateEcriture", "EcrDate")
    cPieceRef  = _getcol(df, "PieceRef", "ReferencePiece", "RefPiece")
    cDebit     = _getcol(df, "Debit")
    cCredit    = _getcol(df, "Credit")
    cLet       = _getcol(df, "EcritureLet", "Lettrage", "CodeLettrage")
    cDateLet   = _getcol(df, "DateLet", "LettrageDate")
    cLib       = _getcol(df, "EcritureLib", "LibelleEcriture")

    # Dates & montants
    cutoff = None
    if cEcrDate:
        df[cEcrDate] = pd.to_datetime(df[cEcrDate], errors="coerce")
        cutoff = pd.Timestamp.utcnow().normalize() - pd.Timedelta(days=aging_days)

    amt_series = _make_amount_series(df, cDebit, cCredit)

    def tier_id(row: pd.Series):
        aux = str(row.get(cCompAuxNum, "") or "").strip() if cCompAuxNum else ""
        auxlib = str(row.get(cCompAuxLib, "") or "").strip() if cCompAuxLib else ""
        if aux:
            return aux, auxlib
        compte = str(row.get(cCompteNum, "") or "").strip() if cCompteNum else ""
        complib = str(row.get(cCompteLib, "") or "").strip() if cCompteLib else ""
        return compte, complib

    # 1) 401/411 non lettrés anciens
    if cCompteNum:
        mask_tiers = _starts_with_series(df[cCompteNum], ("401","411"))
        if cLet or cDateLet:
            m_unlet = _empty_series(df[cLet]) if cLet else pd.Series(True, index=df.index)
            if cDateLet:
                m_unlet = m_unlet | df[cDateLet].isna()
        else:
            m_unlet = pd.Series(True, index=df.index)
        if cEcrDate:
            m_old = df[cEcrDate].notna() & (df[cEcrDate] <= cutoff)
        else:
            m_old = pd.Series(True, index=df.index)
        m_amt = (amt_series >= float(amount_threshold))
        cand = df[mask_tiers & m_unlet & m_old & m_amt].copy()

        rows = []
        for idx, r in cand.iterrows():
            t_id, t_lib = tier_id(r)
            amt = float(amt_series.loc[idx] or 0.0)
            dt = r[cEcrDate] if cEcrDate else None
            rows.append({"tier": t_id or "(inconnu)", "lib": t_lib, "amount": amt, "date": dt})

        if rows:
            tb = pd.DataFrame(rows)
            agg = tb.groupby(["tier","lib"], dropna=False).agg(
                n=("amount","size"),
                total=("amount","sum"),
                oldest=("date","min")
            ).reset_index()
            agg = agg.sort_values("total", ascending=False).head(top_n_per_bucket)
            for _, r in agg.iterrows():
                t = r["tier"]
                lib = r["lib"] or ""
                tot = round(float(r["total"] or 0), 2)
                n = int(r["n"])
                oldest = r["oldest"]
                oldest_str = oldest.date().isoformat() if pd.notna(oldest) else "date inconnue"
                who = "client" if str(t).startswith("411") else "fournisseur"
                label = (
                    f"Postes non lettrés (> {aging_days} j) sur {who} {t} "
                    f"{('('+lib+')' if lib else '')}: {n} écriture(s), total ~{tot}€. "
                    f"Ancienneté depuis {oldest_str}. Indiquez le statut (litige, relance, avoir, plan de règlement)."
                )
                q.append({"type":"texte","question":label})
                q.append({"type":"fichier","question":f"Joindre justificatifs (factures/relevés) pour {t} {('('+lib+')' if lib else '')}."})

    # 2) 401/411 sans référence de pièce
    if cCompteNum and cPieceRef:
        mask_tiers = _starts_with_series(df[cCompteNum], ("401","411"))
        m_nopic = _empty_series(df[cPieceRef])
        m_amt = (amt_series >= float(amount_threshold))
        miss = df[mask_tiers & m_nopic & m_amt].copy()
        if len(miss):
            rows = []
            for idx, r in miss.iterrows():
                t_id, t_lib = tier_id(r)
                amt = float(amt_series.loc[idx] or 0.0)
                rows.append({"tier": t_id or "(inconnu)","lib": t_lib, "amount": amt})
            tb = pd.DataFrame(rows)
            agg = tb.groupby(["tier","lib"]).agg(n=("amount","size"), total=("amount","sum")).reset_index()
            agg = agg.sort_values("total", ascending=False).head(top_n_per_bucket)
            for _, r in agg.iterrows():
                t, lib, n, tot = r["tier"], r["lib"] or "", int(r["n"]), round(float(r["total"] or 0),2)
                q.append({"type":"texte","question":f"{n} écriture(s) sur {t} {('('+lib+')' if lib else '')} sans référence de pièce. Précisez la référence (ou raison)."})
                q.append({"type":"fichier","question":f"Joindre les pièces manquantes pour {t} {('('+lib+')' if lib else '')}."})

    # 3) Doublons de règlements fournisseurs (même date + même montant + même tiers)
    cEcrDate_ok = cEcrDate and (cEcrDate in df.columns)
    if cCompteNum and cEcrDate_ok:
        mask_401 = _starts_with_series(df[cCompteNum], "401")
        tmp = df[mask_401].copy()
        if not tmp.empty:
            tmp[cEcrDate] = pd.to_datetime(tmp[cEcrDate], errors="coerce")
            amounts = _make_amount_series(tmp, cDebit, cCredit)
            tmp["__amt__"] = amounts
            tmp["__date__"] = tmp[cEcrDate].dt.date
            tmp = tmp[tmp["__amt__"] >= float(amount_threshold)]
            tcol = cCompAuxNum or cCompteNum
            if tcol in tmp.columns:
                grp = tmp.groupby([tcol, "__date__", "__amt__"]).size().reset_index(name="n")
                dup = grp[grp["n"] >= 2].head(top_n_per_bucket)
                for _, r in dup.iterrows():
                    t = r[tcol]
                    date = r["__date__"]
                    amt = round(float(r["__amt__"] or 0),2)
                    q.append({"type":"oui_non","question":f"Deux règlements identiques détectés pour le fournisseur {t} le {date} pour ~{amt}€. Confirmez s'il s'agit d'un doublon ?"})
                    q.append({"type":"fichier","question":f"Joindre justificatif (relevé/annulation/avoir) concernant le doublon {t} {date} {amt}€."})

    # 4) Comptes d'attente 471
    if cCompteNum:
        mask_471 = _starts_with_series(df[cCompteNum], "471")
        ca = df[mask_471].copy()
        if len(ca):
            if cLib and (cLib in ca.columns):
                grp = ca.groupby(cLib).size().reset_index(name="n").sort_values("n", ascending=False).head(top_n_per_bucket)
                for _, r in grp.iterrows():
                    lib = str(r[cLib])[:120]
                    q.append({"type":"texte","question":f"Écritures en 471 détectées (ex: '{lib}'). Précisez la nature réelle et le compte définitif."})
                    q.append({"type":"fichier","question":f"Joindre justificatif pour l'écriture en 471 (ex: '{lib}')."})
            else:
                q.append({"type":"texte","question":"Écritures en 471 détectées. Précisez la nature réelle et la régularisation prévue."})
                q.append({"type":"fichier","question":"Joindre justificatifs pour les 471."})

    return q

# -----------------------
# UI
# -----------------------

st.set_page_config(page_title="Form System", page_icon="🧾", layout="centered")

# ✅ Nouveau : API moderne
qp = st.query_params  # dict-like
role = qp.get("role", "client")        # "admin" ou "client"
client_id = qp.get("client", "anonyme")


if role == "admin":
    st.title("👩‍💼 Admin – Création & publication de formulaires")
    pwd = st.text_input("Mot de passe admin", type="password", value="")
    if pwd != ADMIN_PASSWORD:
        st.info("Entrez le mot de passe admin (par défaut: 'admin' – changez la variable d'env ADMIN_PASSWORD).")
        st.stop()

    st.subheader("1) Importer un modèle depuis Excel (optionnel)")
    up = st.file_uploader("Fichier Excel (.xlsx) avec colonnes: type, question, options", type=["xlsx"])

    questions: List[Dict[str, Any]] = st.session_state.get("questions", [])

    if up is not None:
        try:
            df = pd.read_excel(up)
            for col in ["type", "question"]:
                if col not in df.columns:
                    st.error(f"Colonne manquante dans l'Excel: {col}")
                    st.stop()
            if "options" not in df.columns:
                df["options"] = ""

            questions = []
            for _, row in df.iterrows():
                qtype = str(row["type"]).strip()
                if qtype not in SUPPORTED_TYPES:
                    st.warning(f"Type non supporté '{qtype}', question ignorée.")
                    continue
                questions.append({
                    "type": qtype,
                    "question": str(row["question"]).strip(),
                    "options": normalize_options(row.get("options", ""))
                })
            st.session_state["questions"] = questions
            st.success(f"{len(questions)} question(s) importée(s).")
        except Exception as e:
            st.error(f"Erreur de lecture Excel: {e}")
            st.stop()

    st.subheader("2) Éditer / réordonner avant publication")

    # --- 2bis) Génération auto depuis FEC
    st.subheader("2bis) Générer des questions automatiquement depuis un FEC (401/411/471)")
    with st.expander("📥 Importer et analyser un FEC (CSV/TXT/Excel)"):
        fec_file = st.file_uploader("Fichier FEC (CSV/TXT/Excel)", type=["csv","txt","xlsx","xls"], key="fec_upl")
        aging = st.number_input("Ancienneté (jours) pour considérer un poste ancien", min_value=30, max_value=365, value=90, step=15)
        amount_thresh = st.number_input("Seuil de matérialité (montant minimum, €)", min_value=0.0, max_value=1000000.0, value=100.0, step=50.0, format="%.2f")
        topn = st.number_input("Limiter le nombre de questions par catégorie", min_value=5, max_value=200, value=20, step=5)
        if st.button("Analyser le FEC et proposer des questions"):
            if fec_file is None:
                st.warning("Veuillez importer un fichier FEC.")
            else:
                try:
                    df_fec = read_fec_autodetect(fec_file)
                    qs = generate_questions_401_411_471(
                        df_fec,
                        aging_days=int(aging),
                        top_n_per_bucket=int(topn),
                        amount_threshold=float(amount_thresh)
                    )
                    if not qs:
                        st.info("Aucune question pertinente détectée pour 401/411/471 avec ces critères.")
                    else:
                        questions = st.session_state.get("questions", [])
                        questions.extend(qs)
                        st.session_state["questions"] = questions
                        st.success(f"{len(qs)} question(s) générée(s) et ajoutée(s) à l'éditeur. Relisez/modifiez avant publication.")
                except Exception as e:
                    st.error(f"Erreur d'analyse du FEC: {e}")

    if not questions:
        st.info("Aucune question chargée. Importez depuis Excel, générez depuis FEC, ou ajoutez manuellement ci-dessous.")

    # Ajout manuel
    with st.expander("➕ Ajouter une question"):
        new_q_col1, new_q_col2 = st.columns([2,1])
        with new_q_col1:
            new_label = st.text_input("Texte de la question", key="new_label")
        with new_q_col2:
            new_type = st.selectbox("Type", SUPPORTED_TYPES, key="new_type")
        new_opts_str = st.text_input("Options (séparées par ';') pour checkbox_multi", key="new_opts")
        if st.button("Ajouter la question"):
            q = {
                "type": new_type,
                "question": new_label.strip() if new_label else "",
                "options": normalize_options(new_opts_str)
            }
            if q["question"]:
                questions.append(q)
                st.session_state["questions"] = questions
                st.success("Question ajoutée.")
            else:
                st.warning("Le texte de la question est vide.")

    # Edition en liste
    to_delete = []
    reordered = []
    for i, q in enumerate(questions):
        st.markdown(f"**Q{i+1}**")
        c1, c2, c3 = st.columns([3,2,1])
        with c1:
            q["question"] = st.text_input("Intitulé", value=q["question"], key=f"label_{i}")
        with c2:
            q["type"] = st.selectbox("Type", SUPPORTED_TYPES, index=SUPPORTED_TYPES.index(q["type"]), key=f"type_{i}")
        with c3:
            pos = st.number_input("Ordre", min_value=1, max_value=len(questions), value=i+1, key=f"pos_{i}")
        if q["type"] == "checkbox_multi":
            opt_str = "; ".join(q.get("options", []))
            new_opt_str = st.text_input("Options (; séparées)", value=opt_str, key=f"opts_{i}")
            q["options"] = normalize_options(new_opt_str)
        if st.button("Supprimer", key=f"del_{i}"):
            to_delete.append(i)
        reordered.append((pos, q))

    if to_delete:
        for idx in sorted(to_delete, reverse=True):
            questions.pop(idx)
        st.session_state["questions"] = questions
        st.experimental_rerun()

    if st.button("Appliquer l'ordre"):
        reordered.sort(key=lambda x: x[0])
        questions = [q for _, q in reordered]
        st.session_state["questions"] = questions
        st.success("Ordre mis à jour.")

    st.subheader("3) Publication & export")
    client_pub = st.text_input("Identifiant client (pour le lien)", value=query.get("client", [client_id])[0])

    colp1, colp2, colp3 = st.columns(3)
    with colp1:
        if st.button("Publier pour ce client"):
            if not client_pub:
                st.error("Veuillez saisir un identifiant client.")
            elif not questions:
                st.error("Aucune question à publier.")
            else:
                save_form(client_pub, questions)
                st.success(f"Formulaire publié pour '{client_pub}'.")
                st.info(f"Lien client : ?role=client&client={client_pub}")
    with colp2:
        if st.button("Exporter en Word"):
            if not questions:
                st.error("Aucune question à exporter.")
            else:
                path = export_word(questions, client_pub or "client")
                with open(path, "rb") as fh:
                    st.download_button("Télécharger le .docx", fh, file_name=f"{client_pub or 'formulaire'}.docx")
    with colp3:
        if st.button("Vider l'éditeur"):
            st.session_state["questions"] = []
            st.experimental_rerun()

elif role == "client":
    st.title(f"🧾 Formulaire – Client : {client_id}")

    form_def = load_form(client_id)
    if not form_def:
        st.error("Aucun formulaire publié pour ce client.")
        st.stop()

    draft = load_draft(client_id)
    answers = draft.get("answers", {}) if draft else {}

    st.caption("Le brouillon est enregistré automatiquement. Vous pouvez revenir plus tard.")

    data_rows = []
    uploaded_files_map = {}

    for idx, q in enumerate(form_def):
        qkey = f"q_{idx}"
        qtype = q["type"]
        qlabel = q["question"]

        if qtype == "oui_non":
            default = answers.get(qkey, "Non répondu")
            choice = st.radio(qlabel, ["Oui", "Non"], index=(0 if default=="Oui" else 1) if default in ["Oui","Non"] else 1, key=qkey)
            answers[qkey] = choice

        elif qtype == "texte":
            default = answers.get(qkey, "")
            txt = st.text_area(qlabel, value=default, key=qkey)
            answers[qkey] = txt

        elif qtype == "checkbox_multi":
            options = q.get("options", [])
            default = answers.get(qkey, [])
            sels = st.multiselect(qlabel, options, default=default, key=qkey)
            answers[qkey] = sels

        elif qtype == "fichier":
            up = st.file_uploader(qlabel, key=qkey)
            if up is not None:
                client_dir = UPLOADS_DIR / client_id
                client_dir.mkdir(parents=True, exist_ok=True)
                save_path = client_dir / f"{int(time.time())}_{up.name}"
                with open(save_path, "wb") as f:
                    f.write(up.getbuffer())
                answers[qkey] = str(save_path.relative_to(BASE_DIR))
                uploaded_files_map[qkey] = str(save_path)
            else:
                if qkey in answers:
                    st.caption(f"Déjà importé: {answers[qkey]}")

        else:
            st.warning(f"Type non géré: {qtype}")
            continue

        # Autosave brouillon
        save_draft(client_id, {"answers": answers})

    st.divider()

    if st.button("📨 Envoyer"):
        timestamp = pd.Timestamp.utcnow().isoformat()
        for idx, q in enumerate(form_def):
            qkey = f"q_{idx}"
            qlabel = q["question"]
            val = answers.get(qkey, "")
            if isinstance(val, list):
                val = "; ".join(val)
            data_rows.append({
                "timestamp_utc": timestamp,
                "client": client_id,
                "question_index": idx + 1,
                "question": qlabel,
                "type": q["type"],
                "reponse": val
            })
        append_responses_csv(client_id, data_rows)
        save_draft(client_id, {"answers": {}})
        st.success("Merci, vos réponses ont été envoyées.")
        out = RESP_DIR / f"{client_id}.csv"
        if out.exists():
            with open(out, "rb") as fh:
                st.download_button("Télécharger vos réponses (CSV)", fh, file_name=f"{client_id}.csv")
else:
    st.error("Paramètre 'role' invalide. Utilisez ?role=admin ou ?role=client.")
